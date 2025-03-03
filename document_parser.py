import docx
from data_structures import Variable, Module, TestDataManager

class DocumentParser:
    def __init__(self):
        self.data_manager = TestDataManager()
    
    def parse_variable_doc(self, file_path):
        doc = docx.Document(file_path)
        
        print(f"找到 {len(doc.tables)} 个表格")
        
        # 获取文档中的所有表格
        for table in doc.tables:
            print(f"表格行数：{len(table.rows)}")
            # 跳过表头行
            for row in table.rows[1:]:
                # 获取每一行的所有单元格文本
                cells = [cell.text.strip() for cell in row.cells]
                print(f"处理行: {cells}")
                
                if len(cells) >= 9:  # 确保有足够的列
                    try:
                        variable = Variable(
                            name=cells[0],          # 变量名
                            symbol=cells[1],        # 变量符号
                            var_type=cells[2],      # 类型
                            type_desc=cells[3],     # 类型解释
                            initial_value=cells[4],  # 初始值
                            comment=cells[5],       # 变量注释
                            identifier=cells[6],    # 变量标识位
                            min_value=float(cells[7]),  # 最小值
                            max_value=float(cells[8])   # 最大值
                        )
                        self.data_manager.variables[variable.symbol] = variable
                        print(f"成功添加变量: {variable.symbol}")
                    except ValueError as e:
                        print(f"警告：处理行 {cells} 时出错：{e}")
    
    def parse_module_doc(self, file_path):
        doc = docx.Document(file_path)
        current_module = None
        input_vars = []
        output_vars = []
        is_formula = False
        formula_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            if text.startswith('任务名称：'):
                # 如果已经有一个模块，先保存它
                if current_module is not None:
                    current_module.inputs = input_vars
                    current_module.outputs = output_vars
                    if formula_content:
                        current_module.formula = '\n'.join(formula_content)
                    self.data_manager.modules[current_module.name] = current_module
                
                # 创建新模块
                name = text.replace('任务名称：', '').strip()
                current_module = Module(name)
                input_vars = []
                output_vars = []
                is_formula = False
                formula_content = []
                
            elif text.startswith('编号：'):
                is_formula = False
                current_module.number = text.replace('编号：', '').strip()
            elif text.startswith('功能：'):
                is_formula = False
                current_module.function = text.replace('功能：', '').strip()
            elif text.startswith('前置条件：'):
                is_formula = False
                current_module.precondition = text.replace('前置条件：', '').strip()
            elif text.startswith('输入：'):
                is_formula = False
                vars_text = text.replace('输入：', '').strip()
                if vars_text:
                    input_vars = [v.strip() for v in vars_text.split(',')]
            elif text.startswith('输出：'):
                is_formula = False
                vars_text = text.replace('输出：', '').strip()
                if vars_text:
                    output_vars = [v.strip() for v in vars_text.split(',')]
            elif text.startswith('公式：'):
                is_formula = True
                formula_text = text.replace('公式：', '').strip()
                if formula_text:
                    formula_content.append(formula_text)
            elif is_formula:
                # 如果当前是在处理公式部分，继续收集公式内容
                formula_content.append(text)
        
        # 保存最后一个模块
        if current_module is not None:
            current_module.inputs = input_vars
            current_module.outputs = output_vars
            if formula_content:
                current_module.formula = '\n'.join(formula_content)
            self.data_manager.modules[current_module.name] = current_module