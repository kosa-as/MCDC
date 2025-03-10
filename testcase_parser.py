import pandas as pd
import openpyxl
from openpyxl.utils import get_column_letter, range_boundaries
import os
import json
import re
import docx

def unmerge_cells_and_fill(input_file, output_file=None):
    """
    解除Excel文件中的合并单元格，并使每个单元格都显示原来的内容
    
    Args:
        input_file (str): 输入Excel文件路径
        output_file (str, optional): 输出Excel文件路径，默认为在原文件名基础上添加"_unmerged"
    """
    if output_file is None:
        base_name, ext = os.path.splitext(os.path.basename(input_file))
        output_file = os.path.join("output", f"{base_name}_unmerged{ext}")
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    # 加载工作簿
    wb = openpyxl.load_workbook(input_file)
    
    # 处理每个工作表
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        
        # 获取所有合并单元格的范围
        merged_cells = list(ws.merged_cells.ranges)
        
        # 解除合并单元格并填充内容
        for merged_range in merged_cells:
            # 获取合并单元格的范围
            range_str = str(merged_range)
            min_row, min_col, max_row, max_col = range_boundaries(range_str)
            
            # 获取合并单元格的值
            top_left_cell_value = ws.cell(row=min_row, column=min_col).value
            
            # 先解除合并
            ws.unmerge_cells(range_str)
            
            # 再填充每个单元格
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    ws.cell(row=row, column=col).value = top_left_cell_value
        
        print(f"工作表 '{sheet_name}' 中的合并单元格已解除并填充")
    
    # 保存结果
    wb.save(output_file)
    print(f"处理完成，结果已保存至: {output_file}")
    return output_file

def process_input_directory():
    """处理input目录下的所有Excel文件"""
    input_dir = "input"
    
    # 确保input目录存在
    if not os.path.exists(input_dir):
        print(f"错误: 未找到{input_dir}目录")
        return
    
    # 确保output目录存在
    os.makedirs("output", exist_ok=True)
    
    # 获取input目录下的所有Excel文件
    excel_files = [f for f in os.listdir(input_dir) 
                  if f.endswith('.xlsx') or f.endswith('.xls')]
    
    if not excel_files:
        print(f"错误: {input_dir}目录下没有找到Excel文件")
        return
    
    # 处理每个Excel文件
    for excel_file in excel_files:
        input_path = os.path.join(input_dir, excel_file)
        base_name, ext = os.path.splitext(excel_file)
        output_path = os.path.join("output", f"{base_name}_unmerged{ext}")
        
        print(f"正在处理: {input_path}")
        unmerge_cells_and_fill(input_path, output_path)

class Module:
    def __init__(self, name):
        self.name = name
        self.number = ""
        self.function = ""
        self.precondition = ""
        self.inputs = []
        self.outputs = []
        self.formula = ""

class Variable:
    def __init__(self, name="", symbol="", var_type="", type_desc="", initial_value="", 
                 comment="", identifier="", min_value=0, max_value=0):
        self.name = name
        self.symbol = symbol
        self.var_type = var_type
        self.type_desc = type_desc
        self.initial_value = initial_value
        self.comment = comment
        self.identifier = identifier
        self.min_value = min_value
        self.max_value = max_value

def parse_testcase_to_json(input_file, output_file=None):
    """
    从Excel文件中读取测试用例内容并保存为JSON格式
    
    Args:
        input_file (str): 输入Excel文件路径
        output_file (str, optional): 输出JSON文件路径，默认为在原文件名基础上添加"_parsed.json"
    """
    if output_file is None:
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_file = os.path.join("output", f"{base_name}_parsed.json")
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    # 读取Excel文件
    df = pd.read_excel(input_file)
    
    # 存储所有行的数组
    all_rows = []
    
    # 记录上一个非空值
    last_non_empty = {
        "需求编号": "",
        "模块名称": "",
        "前置条件": "",
        "判断条件": ""
    }
    
    # 需要填充的列
    fill_columns = ["需求编号", "模块名称", "前置条件", "判断条件"]
    
    # 逐行处理数据
    for _, row in df.iterrows():
        # 将行数据转换为字典
        row_dict = row.to_dict()
        
        # 将所有值转换为字符串，处理NaN值
        clean_row = {}
        for key, value in row_dict.items():
            key_str = str(key)
            
            if pd.isna(value):
                # 对于需要填充的列，使用上一个非空值
                if key_str in fill_columns:
                    clean_row[key_str] = last_non_empty.get(key_str, "")
                else:
                    clean_row[key_str] = ""
            else:
                value_str = str(value).strip()
                clean_row[key_str] = value_str
                
                # 更新上一个非空值
                if key_str in fill_columns and value_str:
                    last_non_empty[key_str] = value_str
        
        # 添加flag字段，判断条件是否以!开头
        condition = clean_row.get("判断条件", "").strip()
        if condition.startswith("!"):
            clean_row["flag"] = False
            # 去掉条件中的!，保存原始条件
            clean_row["原始判断条件"] = condition
            # 去掉!后的条件
            clean_row["判断条件"] = condition[1:].strip()
        else:
            clean_row["flag"] = True
            clean_row["原始判断条件"] = condition
        
        # 添加到结果数组
        all_rows.append(clean_row)
    
    # 将结果保存为JSON文件
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_rows, f, ensure_ascii=False, indent=4)
    
    print(f"测试用例解析完成，结果已保存至: {output_file}")
    return output_file

def convert_fraction_to_decimal(text):
    """
    将文本中的分数表示(如301/10)转换为小数，并对.1和.9的小数进行就近取整
    
    Args:
        text (str): 包含可能的分数表示的文本
    
    Returns:
        str: 转换后的文本
    """
    # 查找形如数字/数字的模式
    pattern = r'(\-?\d+)/(\d+)'
    
    def replace_fraction(match):
        numerator = int(match.group(1))
        denominator = int(match.group(2))
        
        # 计算小数值
        decimal_value = numerator / denominator
        
        # 检查是否为特殊情况(.1或.9结尾)
        decimal_str = str(decimal_value)
        if decimal_str.endswith('.1'):
            return str(round(decimal_value))
        elif decimal_str.endswith('.9'):
            return str(round(decimal_value))
        else:
            # 保留两位小数，去除尾部的0
            rounded = round(decimal_value, 2)
            return str(rounded).rstrip('0').rstrip('.') if '.' in str(rounded) else str(rounded)
    
    # 替换所有匹配的分数
    return re.sub(pattern, replace_fraction, text)

def merge_testcases_by_requirement(json_file, output_file=None):
    """
    合并相同需求编号和模块名称的测试用例
    
    Args:
        json_file (str): 输入JSON文件路径
        output_file (str, optional): 输出JSON文件路径，默认为在原文件名基础上添加"_merged"
    """
    if output_file is None:
        base_name = os.path.splitext(os.path.basename(json_file))[0]
        output_file = os.path.join("output", f"{base_name}_merged.json")
    
    # 读取JSON文件
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 按需求编号、模块名称和判断条件分组
    grouped_data = {}
    
    for item in data:
        req_id = item.get("需求编号", "")
        module_name = item.get("模块名称", "")
        condition = item.get("判断条件", "").strip()
        flag = item.get("flag", True)  # 默认为True
        testcase = item.get("测试用例", "")
        
        # 跳过没有需求编号或模块名称的项
        if not req_id or not module_name:
            continue
        
        # 创建分组键，包含需求编号、模块名称和判断条件
        key = f"{req_id}_{module_name}_{condition}"
        
        # 如果是新的分组，初始化
        if key not in grouped_data:
            grouped_data[key] = {
                "requirement_id": req_id,
                "module_name": module_name,
                "precondition": item.get("前置条件", ""),
                "condition": condition,
                "true_test_case": [],
                "false_test_case": []
            }
        
        # 如果有测试用例，则添加到相应的数组
        if testcase:
            # 处理测试用例中的分数表示
            processed_testcase = convert_fraction_to_decimal(testcase)
            
            # 根据flag决定添加到哪个数组
            if flag:
                grouped_data[key]["true_test_case"].append(processed_testcase)
            else:
                grouped_data[key]["false_test_case"].append(processed_testcase)
    
    # 将分组数据转换为列表
    result = list(grouped_data.values())
    
    # 将结果保存为JSON文件
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=4)
    
    print(f"测试用例合并完成，结果已保存至: {output_file}")
    return output_file

def clean_formula(formula_text):
    """
    清理公式文本，彻底清除所有空格
    
    Args:
        formula_text (str): 原始公式文本
    
    Returns:
        str: 清理后的公式文本
    """
    # 去除所有空白字符（空格、制表符、换行符等）
    return re.sub(r'\s+', '', formula_text)

def parse_module_doc(input_file, output_file=None):
    """
    从Word文档中读取模块信息并保存为JSON格式
    
    Args:
        input_file (str): 输入Word文档路径
        output_file (str, optional): 输出JSON文件路径，默认为在原文件名基础上添加"_modules.json"
    """
    if output_file is None:
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_file = os.path.join("output", f"{base_name}_modules.json")
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    # 读取Word文档
    doc = docx.Document(input_file)
    
    # 存储模块信息
    modules = []
    current_module = None
    input_vars = []
    output_vars = []
    is_formula = False
    formula_content = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        if text.startswith('任务名称：'):
            # 如果已经有一个模块，先保存它
            if current_module is not None:
                # 处理公式内容，彻底清除所有空格
                formula_text = ""
                if formula_content:
                    # 将所有公式内容合并为一个字符串，然后清除所有空格
                    combined_formula = ''.join(formula_content)
                    formula_text = clean_formula(combined_formula)
                
                module_dict = {
                    "name": current_module.name,
                    "number": current_module.number,
                    "function": current_module.function,
                    "precondition": current_module.precondition,
                    "inputs": input_vars,
                    "outputs": output_vars,
                    "formula": formula_text
                }
                modules.append(module_dict)
            
            # 创建新模块
            name = text.replace('任务名称：', '').strip()
            current_module = Module(name)
            input_vars = []
            output_vars = []
            is_formula = False
            formula_content = []
            
        elif text.startswith('编号：'):
            is_formula = False
            current_module.number = text.replace('编号：', '').strip()
        elif text.startswith('功能：'):
            is_formula = False
            current_module.function = text.replace('功能：', '').strip()
        elif text.startswith('前置条件：'):
            is_formula = False
            current_module.precondition = text.replace('前置条件：', '').strip()
        elif text.startswith('输入：'):
            is_formula = False
            vars_text = text.replace('输入：', '').strip()
            if vars_text:
                # 分割并清理每个变量名
                input_vars = [v.strip() for v in vars_text.split(',')]
        elif text.startswith('输出：'):
            is_formula = False
            vars_text = text.replace('输出：', '').strip()
            if vars_text:
                # 分割并清理每个变量名
                output_vars = [v.strip() for v in vars_text.split(',')]
        elif text.startswith('公式：'):
            is_formula = True
            formula_text = text.replace('公式：', '').strip()
            if formula_text:
                formula_content.append(formula_text)
        elif is_formula:
            # 如果当前是在处理公式部分，继续收集公式内容
            formula_content.append(text)
    
    # 保存最后一个模块
    if current_module is not None:
        # 处理公式内容，彻底清除所有空格
        formula_text = ""
        if formula_content:
            # 将所有公式内容合并为一个字符串，然后清除所有空格
            combined_formula = ''.join(formula_content)
            formula_text = clean_formula(combined_formula)
        
        module_dict = {
            "name": current_module.name,
            "number": current_module.number,
            "function": current_module.function,
            "precondition": current_module.precondition,
            "inputs": input_vars,
            "outputs": output_vars,
            "formula": formula_text
        }
        modules.append(module_dict)
    
    # 将结果保存为JSON文件
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(modules, f, ensure_ascii=False, indent=4)
    
    print(f"模块信息解析完成，结果已保存至: {output_file}")
    return output_file

def extract_code_blocks(formula, condition):
    """
    从公式中提取与条件相关的代码块
    
    Args:
        formula (str): 公式内容
        condition (str): 条件表达式
    
    Returns:
        tuple: (true_result, false_result) 条件为真和为假时的代码块
    """
    print(f"\n调试 - 提取代码块:")
    print(f"条件: '{condition}'")
    print(f"公式长度: {len(formula)}")
    print(f"公式前100个字符: '{formula[:100]}...'")
    
    # 如果条件为空，直接返回空结果
    if not condition:
        print("条件为空，返回空结果")
        return "", ""
    
    # 直接在公式中查找条件
    condition_pos = formula.find(condition)
    print(f"条件在公式中的位置: {condition_pos}")
    
    if condition_pos == -1:
        print("未找到条件，尝试查找if语句")
        # 尝试查找if语句
        if_pattern = f"if({condition})"
        if_pos = formula.find(if_pattern)
        print(f"if语句在公式中的位置: {if_pos}")
        
        if if_pos != -1:
            condition_pos = if_pos
        else:
            print("未找到条件，返回空结果")
            return "", ""
    
    # 从条件位置开始查找第一个左花括号
    brace_start = formula.find('{', condition_pos)
    print(f"左花括号位置: {brace_start}")
    
    if brace_start == -1:
        print("未找到左花括号，返回空结果")
        return "", ""
    
    # 查找匹配的右花括号
    brace_count = 1
    brace_end = brace_start + 1
    
    while brace_count > 0 and brace_end < len(formula):
        if formula[brace_end] == '{':
            brace_count += 1
        elif formula[brace_end] == '}':
            brace_count -= 1
        brace_end += 1
    
    print(f"右花括号位置: {brace_end}")
    
    if brace_count != 0:
        print("花括号不匹配，返回空结果")
        return "", ""  # 花括号不匹配
    
    # 提取true_result
    true_result = formula[brace_start+1:brace_end-1]
    print(f"提取的true_result: '{true_result}'")
    
    # 查找else关键字
    else_pos = formula.find('else', brace_end)
    print(f"else关键字位置: {else_pos}")
    
    false_result = ""
    
    if else_pos != -1 and else_pos < brace_end + 20:  # 确保else在右花括号附近
        print(f"找到else关键字，在右花括号附近")
        # 查找else后的第一个左花括号
        else_brace_start = formula.find('{', else_pos)
        print(f"else后的左花括号位置: {else_brace_start}")
        
        if else_brace_start != -1:
            # 查找匹配的右花括号
            else_brace_count = 1
            else_brace_end = else_brace_start + 1
            
            while else_brace_count > 0 and else_brace_end < len(formula):
                if formula[else_brace_end] == '{':
                    else_brace_count += 1
                elif formula[else_brace_end] == '}':
                    else_brace_count -= 1
                else_brace_end += 1
            
            print(f"else后的右花括号位置: {else_brace_end}")
            
            if else_brace_count == 0:
                false_result = formula[else_brace_start+1:else_brace_end-1]
                print(f"提取的false_result: '{false_result}'")
            else:
                print("else后的花括号不匹配")
        else:
            print("未找到else后的左花括号")
    else:
        print("未找到else关键字或else不在右花括号附近")
    
    return true_result, false_result

def match_and_extract_results(testcase_file, module_file, output_file=None):
    """
    匹配测试用例和模块信息，提取结果
    
    Args:
        testcase_file (str): 测试用例JSON文件路径
        module_file (str): 模块信息JSON文件路径
        output_file (str, optional): 输出JSON文件路径
    """
    if output_file is None:
        base_name = os.path.splitext(os.path.basename(testcase_file))[0]
        output_file = os.path.join("output", f"{base_name}_with_results.json")
    
    # 读取测试用例文件
    with open(testcase_file, 'r', encoding='utf-8') as f:
        testcases = json.load(f)
    
    # 读取模块信息文件
    with open(module_file, 'r', encoding='utf-8') as f:
        modules = json.load(f)
    
    # 创建模块查找索引
    module_index = {}
    for module in modules:
        # 使用number和name作为键
        key = f"{module['number']}_{module['name']}"
        module_index[key] = module
    
    # 中文标点符号替换为英文标点符号的映射
    punctuation_map = {
        '（': '(',
        '）': ')',
        '，': ',',
        '。': '.',
        '：': ':',
        '；': ';',
        '"': '"',
        '"': '"',
        ''': "'",
        ''': "'",
        '【': '[',
        '】': ']',
        '《': '<',
        '》': '>',
        '！': '!',
        '？': '?',
        '、': ',',
        '…': '...'
    }
    
    # 遍历测试用例，匹配模块并提取结果
    for testcase in testcases:
        req_id = testcase.get("requirement_id", "")
        module_name = testcase.get("module_name", "").strip('"')
        condition = testcase.get("condition", "")
        
        print(f"\n处理测试用例: {req_id} - {module_name}")
        
        # 替换中文标点符号
        for cn_punct, en_punct in punctuation_map.items():
            req_id = req_id.replace(cn_punct, en_punct)
            module_name = module_name.replace(cn_punct, en_punct)
            condition = condition.replace(cn_punct, en_punct)
        
        # 更新测试用例中的字段
        testcase["requirement_id"] = req_id
        testcase["module_name"] = module_name
        
        # 清理条件，去掉所有空格和外层括号
        if condition:
            # 去掉所有空格
            condition = re.sub(r'\s+', '', condition)
            # 如果条件被括号包围，去掉外层括号
            if condition.startswith('(') and condition.endswith(')'):
                condition = condition[1:-1]
            # 更新测试用例中的条件
            testcase["condition"] = condition
            print(f"清理后的条件: '{condition}'")
        
        # 初始化结果字段
        testcase["true_result"] = ""
        testcase["false_result"] = ""
        
        # 查找匹配的模块
        found_module = False
        for module in modules:
            if module['number'] == req_id and module['name'] == module_name:
                found_module = True
                formula = module.get("formula", "")
                print(f"找到匹配的模块: {module['number']} - {module['name']}")
                print(f"公式长度: {len(formula)}")
                
                # 提取结果
                true_result, false_result = extract_code_blocks(formula, condition)
                testcase["true_result"] = true_result
                testcase["false_result"] = false_result
                break
        
        if not found_module:
            print(f"未找到匹配的模块")
    
    # 保存结果
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(testcases, f, ensure_ascii=False, indent=4)
    
    print(f"结果提取完成，已保存至: {output_file}")
    return output_file

def parse_data_doc(input_file, output_file=None):
    """
    从Word文档中提取变量和常量信息，并将常量保存为JSON格式
    
    Args:
        input_file (str): 输入Word文档路径
        output_file (str, optional): 输出JSON文件路径，默认为在原文件名基础上添加"_constants.json"
    """
    if output_file is None:
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_file = os.path.join("output", f"{base_name}_constants.json")
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    # 读取Word文档
    doc = docx.Document(input_file)
    
    # 存储变量和常量
    variables = {}
    constants = {}
    
    print(f"找到 {len(doc.tables)} 个表格")
    
    # 获取文档中的所有表格
    for table_index, table in enumerate(doc.tables):
        print(f"处理表格 {table_index+1}，行数：{len(table.rows)}")
        
        # 跳过表头行
        for row_index, row in enumerate(table.rows[1:], 1):
            # 获取每一行的所有单元格文本
            cells = [cell.text.strip() for cell in row.cells]
            print(f"处理行 {row_index}: {cells[:2]}")  # 只打印前两列，避免输出过多
            
            if len(cells) >= 9:  # 确保有足够的列
                try:
                    symbol = cells[1]
                    name = cells[0]
                    var_type = cells[2]
                    
                    # 尝试转换初始值、最小值和最大值为浮点数
                    try:
                        initial_val_str = cells[4].strip()
                        initial_val = float(initial_val_str) if initial_val_str else None
                    except ValueError:
                        initial_val = None
                        
                    try:
                        min_val_str = cells[7].strip()
                        min_val = float(min_val_str) if min_val_str else None
                    except ValueError:
                        min_val = None
                        
                    try:
                        max_val_str = cells[8].strip()
                        max_val = float(max_val_str) if max_val_str else None
                    except ValueError:
                        max_val = None
                    
                    # 判断是否为常量
                    is_constant = False
                    constant_value = None
                    
                    # 特殊处理level_default
                    if symbol == "level_default":
                        is_constant = True
                        constant_value = 0
                        print(f"特殊处理常量: {symbol} = {constant_value}")
                    # 判断最小值等于最大值等于初始值的情况
                    elif min_val is not None and max_val is not None and min_val == max_val:
                        if initial_val is not None:
                            if min_val == initial_val:
                                is_constant = True
                                constant_value = min_val
                        else:
                            is_constant = True
                            constant_value = min_val
                    # 判断只有初始值的情况
                    elif min_val is None and max_val is None and initial_val is not None:
                        is_constant = True
                        constant_value = initial_val
                    
                    if is_constant and constant_value is not None:
                        # 将浮点数转换为整数（如果可以整除1）
                        if isinstance(constant_value, float) and constant_value.is_integer():
                            constants[symbol] = int(constant_value)
                        else:
                            constants[symbol] = constant_value
                        print(f"添加常量: {symbol} = {constants[symbol]}")
                    else:
                        variable = Variable(
                            name=name,
                            symbol=symbol,
                            var_type=var_type,
                            type_desc=cells[3],
                            initial_value=cells[4],
                            comment=cells[5],
                            identifier=cells[6],
                            min_value=min_val if min_val is not None else 0,
                            max_value=max_val if max_val is not None else 0
                        )
                        variables[symbol] = {
                            "name": variable.name,
                            "symbol": variable.symbol,
                            "var_type": variable.var_type,
                            "type_desc": variable.type_desc,
                            "initial_value": variable.initial_value,
                            "comment": variable.comment,
                            "identifier": variable.identifier,
                            "min_value": variable.min_value,
                            "max_value": variable.max_value
                        }
                        print(f"添加变量: {symbol}")
                except Exception as e:
                    print(f"警告：处理行时出错：{e}")
    
    # 将常量保存为JSON文件
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(constants, f, ensure_ascii=False, indent=4)
    
    print(f"常量提取完成，共 {len(constants)} 个，结果已保存至: {output_file}")
    return output_file

if __name__ == "__main__":
    # 确保输出目录存在
    os.makedirs("output", exist_ok=True)
    
    # 解析Excel文件到JSON
    input_excel = "input/testcase.xlsx"
    parsed_json_file = "output/testcase_parsed.json"
    merged_json_file = "output/testcase_parsed_merged.json"
    
    print(f"正在解析测试用例文件: {input_excel}")
    parse_testcase_to_json(input_excel, parsed_json_file)
    
    # 合并相同需求编号和模块名称的测试用例
    print(f"正在合并测试用例...")
    merge_testcases_by_requirement(parsed_json_file)
    
    # 解析模块文档
    input_doc = "input/module.docx"
    module_json_file = "output/module_modules.json"
    
    if os.path.exists(input_doc):
        print(f"正在解析模块文档: {input_doc}")
        parse_module_doc(input_doc)
        
        # 匹配测试用例和模块信息，提取结果
        if os.path.exists(merged_json_file) and os.path.exists(module_json_file):
            print(f"正在匹配测试用例和模块信息，提取结果...")
            match_and_extract_results(merged_json_file, module_json_file)
    else:
        print(f"未找到模块文档: {input_doc}")
    
    # 解析数据文档
    input_data_doc = "input/data.docx"
    if os.path.exists(input_data_doc):
        print(f"正在解析数据文档: {input_data_doc}")
        parse_data_doc(input_data_doc)
    else:
        print(f"未找到数据文档: {input_data_doc}")
